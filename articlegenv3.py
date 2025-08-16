#!/usr/bin/env python3
"""
Comprehensive Research Article Generator - IMPROVED VERSION
Automatically generates research articles from topic input with proper citations and APA7 formatting.

Key Improvements:
1. Fixed paper filtering issues that were removing all papers
2. Added fallback mechanisms for empty results
3. Improved error handling and logging
4. Enhanced search quality and relevance scoring
5. Better citation management and formatting
6. More robust content generation with retry logic
"""

import os
import sys
import json
import time
import logging
import argparse
import requests
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, asdict
from datetime import datetime, timedelta
import re
import statistics

# Core libraries
import openai
from scholarly import scholarly
import arxiv
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from textstat import flesch_reading_ease

# Document processing
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown2

# Citation and bibliography
from pybtex.database import BibliographyData, Entry
from pybtex.style.formatting.alpha import Style
from pybtex.backends.latex import Backend

# Configuration and utilities
import yaml
from tqdm import tqdm

# Setup logging with better formatting
logging.basicConfig(
    level=logging.INFO, 
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler('article_generator.log')
    ]
)
logger = logging.getLogger(__name__)

@dataclass
class ResearchPaper:
    """Enhanced data structure for a research paper"""
    title: str
    authors: List[str]
    year: int
    abstract: str
    url: str
    doi: str = ""
    venue: str = ""
    citations: int = 0
    key_findings: List[str] = None
    relevance_score: float = 0.0
    quality_score: float = 0.0
    source: str = ""  # Track which database this came from
    
    def __post_init__(self):
        if self.key_findings is None:
            self.key_findings = []
        
        # Calculate quality score based on multiple factors
        self.quality_score = self._calculate_quality_score()
    
    def _calculate_quality_score(self) -> float:
        """Calculate a quality score for the paper"""
        score = 0.0
        
        # Title quality (not too short, not too long)
        title_len = len(self.title.split()) if self.title else 0
        if 5 <= title_len <= 20:
            score += 1.0
        elif title_len > 0:
            score += 0.5
        
        # Abstract quality
        abstract_len = len(self.abstract.split()) if self.abstract else 0
        if abstract_len >= 50:
            score += 2.0
        elif abstract_len >= 20:
            score += 1.0
        
        # Citation score (normalized)
        if self.citations > 0:
            score += min(2.0, self.citations / 50.0)  # Max 2 points for citations
        
        # Venue quality
        if self.venue:
            score += 0.5
        
        # DOI presence
        if self.doi:
            score += 0.5
        
        # Recency bonus (papers from last 5 years get bonus)
        current_year = datetime.now().year
        if current_year - self.year <= 5:
            score += 1.0
        elif current_year - self.year <= 10:
            score += 0.5
        
        return score

@dataclass
class ArticleSection:
    """Enhanced data structure for an article section"""
    title: str
    content: str
    word_count: int = 0
    citations: List[str] = None
    quality_score: float = 0.0
    
    def __post_init__(self):
        if self.citations is None:
            self.citations = []
        if self.word_count == 0:
            self.word_count = len(self.content.split())
        
        # Calculate quality score
        self.quality_score = self._calculate_quality_score()
    
    def _calculate_quality_score(self) -> float:
        """Calculate quality score for the section"""
        score = 0.0
        
        # Length appropriateness
        if 100 <= self.word_count <= 2000:
            score += 2.0
        elif self.word_count > 0:
            score += 1.0
        
        # Citation density (for academic sections)
        if self.word_count > 0:
            citation_density = len(self.citations) / (self.word_count / 100)
            score += min(1.0, citation_density)
        
        # Content structure (basic check for paragraphs)
        paragraph_count = self.content.count('\n\n') + 1
        if paragraph_count >= 2:
            score += 1.0
        
        return score

class Config:
    """Enhanced configuration management with validation"""
    def __init__(self, config_path: str = "config.yaml"):
        self.config_path = config_path
        self.config = self._load_config()
        self._validate_config()
    
    def _load_config(self) -> Dict[str, Any]:
        """Load configuration from YAML file with better defaults"""
        default_config = {
            "apis": {
                "openai_api_key": os.getenv("OPENAI_API_KEY", ""),
                "semantic_scholar_api_key": os.getenv("SEMANTIC_SCHOLAR_API_KEY", ""),
                "quillbot_api_key": os.getenv("QUILLBOT_API_KEY", ""),
                "scispace_api_key": os.getenv("SCISPACE_API_KEY", "")
            },
            "search": {
                "max_papers": 25,  # Increased from 20
                "min_citation_count": 0,  # Reduced from 5 to be more inclusive
                "max_year_range": 15,  # Increased from 10
                "search_sources": ["semantic_scholar", "arxiv"],
                "quality_threshold": 2.0,  # Minimum quality score
                "min_abstract_length": 50,  # Minimum words in abstract
                "max_results_per_source": 15
            },
            "generation": {
                "model": "gpt-5-mini", #changed from gpt-4
                "temperature": 1.0, #1.0,
                "max_completion_tokens": 3500,  # Increased from 2000
                "retry_attempts": 1,
                "target_word_counts": {
                    "abstract": 250,
                    "introduction": 400,
                    "literature_review": 1000,
                    "method": 600,
                    "results": 800,
                    "conclusion": 400
                },
                "fallback_model": "gpt-4-turbo"
            },
            "output": {
                "format": ["docx", "markdown","pdf"],
                "output_dir": "outputs",
                "template_path": "templates/apa7_template.docx",
                "include_summary": True
            },
            "quality": {
                "min_section_words": 100,
                "max_section_words": 2500,
                "target_readability": 40,  # Flesch Reading Ease score
                "require_citations": True
            }
        }
        
        if os.path.exists(self.config_path):
            try:
                with open(self.config_path, 'r') as f:
                    user_config = yaml.safe_load(f)
                    if user_config:
                        default_config.update(user_config)
            except Exception as e:
                import sys
                lineno = sys.exc_info()[2].tb_lineno
                logger.warning(f"Error loading config file {self.config_path} at line {lineno}: {e}")
                logger.info("Using default configuration")
        else:
            # Create default config file
            try:
                with open(self.config_path, 'w') as f:
                    yaml.dump(default_config, f, default_flow_style=False, indent=2)
                logger.info(f"Created default config file: {self.config_path}")
            except Exception as e:
                import sys
                lineno = sys.exc_info()[2].tb_lineno
                logger.warning(f"Could not create config file at line {lineno}: {e}")
        
        return default_config
    
    def _validate_config(self):
        """Validate configuration values"""
        # Check required API keys
        if not self.get("apis.openai_api_key"):
            logger.warning("OpenAI API key not found. Set OPENAI_API_KEY environment variable.")
        
        # Validate search parameters
        max_papers = self.get("search.max_papers", 25)
        if max_papers < 5:
            logger.warning("max_papers is very low, may result in insufficient data")
        
        # Validate model
        model = self.get("generation.model", "gpt-5-mini")
        if model not in ["gpt-5","gpt-5-mini","gpt-4", "gpt-3.5-turbo", "gpt-4-turbo"]:
            logger.warning(f"Unknown model: {model}")
    
    def get(self, key: str, default=None):
        """Get configuration value using dot notation"""
        keys = key.split('.')
        value = self.config
        for k in keys:
            if isinstance(value, dict):
                value = value.get(k, default)
            else:
                return default
            if value is None:
                return default
        return value

class TopicRefiner:
    """Enhanced topic refinement with better academic formatting"""
    
    @staticmethod
    def refine_topic(topic: str) -> Dict[str, str]:
        """Convert a topic into research title and question with better refinement"""
        topic = topic.strip()
        
        # Academic keywords for title generation
        academic_keywords = {
            "impact": "Impact of",
            "effect": "Effects of", 
            "analysis": "Analysis of",
            "study": "A Study on",
            "review": "A Review of",
            "comparison": "Comparative Analysis of",
            "evaluation": "Evaluation of",
            "assessment": "Assessment of",
            "investigation": "Investigation into",
            "exploration": "Exploring",
            "relationship": "The Relationship between",
            "role": "The Role of",
            "influence": "The Influence of"
        }
        
        # Clean and normalize topic
        topic_lower = topic.lower()
        
        # Generate research title
        title = topic.title()
        if not any(keyword in topic_lower for keyword in academic_keywords.keys()):
            # Add appropriate academic framing
            if "and" in topic_lower or "vs" in topic_lower:
                title = f"Comparative Analysis of {title}"
            elif any(word in topic_lower for word in ["how", "why", "what", "when", "where"]):
                title = f"An Investigation into {title}"
            else:
                title = f"A Comprehensive Analysis of {title}"
        
        # Generate multiple research questions
        questions = [
            f"What are the key aspects and implications of {topic}?",
            f"How does {topic} impact current research and practice?",
            f"What are the main findings in recent literature regarding {topic}?"
        ]
        
        # Generate enhanced search terms
        search_terms = TopicRefiner._extract_search_terms(topic)
        
        return {
            "title": title,
            "research_question": questions[0],
            "alternative_questions": questions[1:],
            "search_terms": search_terms,
            "original_topic": topic
        }
    
    @staticmethod
    def _extract_search_terms(topic: str) -> List[str]:
        """Extract meaningful search terms from topic"""
        # Remove common stopwords and short words
        stopwords_set = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by'}
        
        # Split and clean
        words = re.findall(r'\b\w+\b', topic.lower())
        search_terms = [word for word in words if len(word) > 2 and word not in stopwords_set]
        
        # Add some related academic terms
        academic_terms = ['research', 'study', 'analysis', 'review']
        search_terms.extend([term for term in academic_terms if term not in search_terms])
        
        return search_terms[:8]  # Limit to 8 terms

class PaperSearcher:
    """Enhanced paper searcher with better filtering and error handling"""
    
    def __init__(self, config: Config):
        self.config = config
        self.papers = []
        self.search_stats = {
            "total_found": 0,
            "after_deduplication": 0,
            "after_filtering": 0,
            "by_source": {}
        }
    
    def search_semantic_scholar(self, query: str, limit: int = 15) -> List[ResearchPaper]:
        """Enhanced Semantic Scholar search with better error handling"""
        papers = []
        try:
            base_url = "https://api.semanticscholar.org/graph/v1/paper/search"
            params = {
                "query": query,
                "limit": limit,
                "fields": "title,authors,year,abstract,url,citationCount,venue,externalIds,fieldsOfStudy"
            }
            
            headers = {"User-Agent": "ResearchArticleGenerator/1.0"}
            api_key = self.config.get("apis.semantic_scholar_api_key")
            if api_key:
                headers["x-api-key"] = api_key
            
            response = requests.get(base_url, params=params, headers=headers, timeout=30)
            response.raise_for_status()
            
            data = response.json()
            
            for paper_data in data.get("data", []):
                try:
                    # More lenient abstract requirement
                    abstract = paper_data.get("abstract", "")
                    if abstract and len(abstract.split()) >= self.config.get("search.min_abstract_length", 50):
                        paper = ResearchPaper(
                            title=paper_data.get("title", "").strip(),
                            authors=[author.get("name", "") for author in paper_data.get("authors", [])],
                            year=paper_data.get("year") or datetime.now().year,
                            abstract=abstract,
                            url=paper_data.get("url", ""),
                            venue=paper_data.get("venue", ""),
                            citations=paper_data.get("citationCount", 0),
                            doi=paper_data.get("externalIds", {}).get("DOI", ""),
                            source="Semantic Scholar"
                        )
                        
                        # Calculate relevance score
                        paper.relevance_score = self._calculate_relevance(paper, query)
                        papers.append(paper)
                except Exception as e:
                    import sys
                    lineno = sys.exc_info()[2].tb_lineno
                    logger.warning(f"Error processing Semantic Scholar paper at line {lineno}: {e}")
                    continue
            
            logger.info(f"Semantic Scholar: Found {len(papers)} valid papers")
            
        except requests.exceptions.RequestException as e:
            import sys
            lineno = sys.exc_info()[2].tb_lineno
            logger.error(f"Semantic Scholar API request failed at line {lineno}: {e}")
        except Exception as e:
            import sys
            lineno = sys.exc_info()[2].tb_lineno
            logger.error(f"Semantic Scholar search error at line {lineno}: {e}")
        
        return papers
    
    def search_google_scholar(self, query: str, limit: int = 10) -> List[ResearchPaper]:
        """Enhanced Google Scholar search with better error handling"""
        papers = []
        try:
            search_query = scholarly.search_pubs(query)
            
            count = 0
            for result in search_query:
                if count >= limit:
                    break
                
                try:
                    # Fill in more details with timeout protection
                    filled_result = scholarly.fill(result)
                    
                    # Extract abstract
                    abstract = filled_result.get("bib", {}).get("abstract", "")
                    if not abstract and "eprint" in filled_result:
                        abstract = filled_result.get("eprint", "")[:500]  # Fallback to partial text
                    
                    if abstract and len(abstract.split()) >= 20:  # More lenient requirement
                        year = filled_result.get("bib", {}).get("pub_year")
                        if isinstance(year, str):
                            year = int(year) if year.isdigit() else datetime.now().year
                        elif not year:
                            year = datetime.now().year
                        
                        paper = ResearchPaper(
                            title=filled_result.get("bib", {}).get("title", "").strip(),
                            authors=filled_result.get("bib", {}).get("author", []),
                            year=year,
                            abstract=abstract,
                            url=filled_result.get("pub_url", ""),
                            venue=filled_result.get("bib", {}).get("venue", ""),
                            citations=filled_result.get("num_citations", 0),
                            source="Google Scholar"
                        )
                        
                        paper.relevance_score = self._calculate_relevance(paper, query)
                        papers.append(paper)
                        count += 1
                
                except Exception as e:
                    import sys
                    lineno = sys.exc_info()[2].tb_lineno
                    logger.warning(f"Error processing Google Scholar result at line {lineno}: {e}")
                    continue
                    
                time.sleep(2)  # Be respectful to Google Scholar
            
            logger.info(f"Google Scholar: Found {len(papers)} valid papers")
                
        except Exception as e:
            import sys
            lineno = sys.exc_info()[2].tb_lineno
            logger.error(f"Google Scholar search error at line {lineno}: {e}")
        
        return papers
    
    def search_arxiv(self, query: str, limit: int = 15) -> List[ResearchPaper]:
        """Enhanced arXiv search"""
        papers = []
        try:
            client = arxiv.Client()
            search = arxiv.Search(
                query=query,
                max_results=limit,
                sort_by=arxiv.SortCriterion.Relevance
            )
            
            for result in client.results(search):
                paper = ResearchPaper(
                    title=result.title.strip(),
                    authors=[author.name for author in result.authors],
                    year=result.published.year,
                    abstract=result.summary,
                    url=result.pdf_url,
                    doi=result.doi or "",
                    venue="arXiv",
                    source="arXiv"
                )
                
                paper.relevance_score = self._calculate_relevance(paper, query)
                papers.append(paper)
            
            logger.info(f"arXiv: Found {len(papers)} papers")
                
        except Exception as e:
            import sys
            lineno = sys.exc_info()[2].tb_lineno
            logger.error(f"arXiv search error at line {lineno}: {e}")
        
        return papers
    
    def _calculate_relevance(self, paper: ResearchPaper, query: str) -> float:
        """Calculate relevance score for a paper based on query"""
        score = 0.0
        query_terms = set(query.lower().split())
        
        # Title relevance
        title_terms = set(paper.title.lower().split()) if paper.title else set()
        title_overlap = len(query_terms.intersection(title_terms))
        score += title_overlap * 2.0
        
        # Abstract relevance
        if paper.abstract:
            abstract_terms = set(paper.abstract.lower().split())
            abstract_overlap = len(query_terms.intersection(abstract_terms))
            score += abstract_overlap * 0.5
        
        # Boost for exact phrase matches
        query_lower = query.lower()
        if paper.title and query_lower in paper.title.lower():
            score += 5.0
        if paper.abstract and query_lower in paper.abstract.lower():
            score += 2.0
        
        return score
    
    def search_all_sources(self, query: str) -> List[ResearchPaper]:
        """Enhanced search with better source management and fallbacks"""
        all_papers = []
        sources = self.config.get("search.search_sources", ["semantic_scholar", "arxiv"])
        max_per_source = self.config.get("search.max_results_per_source", 15)
        
        logger.info(f"Searching for papers on: '{query}'")
        logger.info(f"Using sources: {sources}")
        
        # Search each source
        for source in sources:
            try:
                papers = []
                if source == "semantic_scholar":
                    papers = self.search_semantic_scholar(query, max_per_source)
                elif source == "google_scholar":
                    pass #Giving captcha error
                    # papers = self.search_google_scholar(query, max_per_source)
                elif source == "arxiv":
                    papers = self.search_arxiv(query, max_per_source)
                
                all_papers.extend(papers)
                self.search_stats["by_source"][source] = len(papers)
                
            except Exception as e:
                import sys
                lineno = sys.exc_info()[2].tb_lineno
                logger.error(f"Error searching {source} at line {lineno}: {e}")
                continue
        
        self.search_stats["total_found"] = len(all_papers)
        logger.info(f"Total papers found across all sources: {len(all_papers)}")
        
        if not all_papers:
            logger.warning("No papers found from any source!")
            # Try a broader search with modified query
            broader_query = " ".join(query.split()[:3])  # Use fewer terms
            logger.info(f"Trying broader search with: '{broader_query}'")
            
            for source in sources:
                try:
                    if source == "semantic_scholar":
                        papers = self.search_semantic_scholar(broader_query, max_per_source * 2)
                    elif source == "arxiv":
                        papers = self.search_arxiv(broader_query, max_per_source * 2)
                    all_papers.extend(papers)
                except Exception as e:
                    import sys
                    lineno = sys.exc_info()[2].tb_lineno
                    logger.error(f"Broader search failed for {source} at line {lineno}: {e}")
        
        # Filter and deduplicate with improved logic
        filtered_papers = self._filter_and_deduplicate(all_papers)
        self.search_stats["after_deduplication"] = len(set(p.title.lower() for p in all_papers if p.title))
        self.search_stats["after_filtering"] = len(filtered_papers)
        
        logger.info(f"Papers after filtering and deduplication: {len(filtered_papers)}")
        
        return filtered_papers
    
    def _filter_and_deduplicate(self, papers: List[ResearchPaper]) -> List[ResearchPaper]:
        """Improved filtering with more lenient criteria and better fallbacks"""
        if not papers:
            return []
        
        seen_titles = set()
        filtered_papers = []
        current_year = datetime.now().year
        
        # Configuration
        min_citations = self.config.get("search.min_citation_count", 0)
        max_age = self.config.get("search.max_year_range", 15)
        quality_threshold = self.config.get("search.quality_threshold", 2.0)
        
        logger.info(f"Filtering criteria:")
        logger.info(f"  - Min citations: {min_citations}")
        logger.info(f"  - Max age: {max_age} years")
        logger.info(f"  - Quality threshold: {quality_threshold}")
        
        # First pass: strict filtering
        for paper in papers:
            # Skip if no title or abstract
            if not paper.title or not paper.abstract:
                continue
            
            # Normalize title for comparison
            title_normalized = re.sub(r'[^\w\s]', '', paper.title.lower().strip())
            if title_normalized in seen_titles:
                continue
            seen_titles.add(title_normalized)
            
            # Apply filters
            age = current_year - paper.year
            if (paper.citations >= min_citations and 
                age <= max_age and
                paper.quality_score >= quality_threshold):
                filtered_papers.append(paper)
        
        logger.info(f"After strict filtering: {len(filtered_papers)} papers")
        
        # If we have too few papers, use more lenient criteria
        if len(filtered_papers) < 10:
            logger.info("Too few papers, applying lenient filtering...")
            filtered_papers = []
            seen_titles = set()
            
            # Sort by quality and relevance
            papers_sorted = sorted(papers, key=lambda p: (p.quality_score + p.relevance_score), reverse=True)
            
            for paper in papers_sorted:
                if len(filtered_papers) >= 20:  # Stop at reasonable number
                    break
                
                if not paper.title or not paper.abstract:
                    continue
                
                title_normalized = re.sub(r'[^\w\s]', '', paper.title.lower().strip())
                if title_normalized in seen_titles:
                    continue
                seen_titles.add(title_normalized)
                
                # More lenient criteria
                age = current_year - paper.year
                if (age <= 20 and  # Allow older papers
                    len(paper.abstract.split()) >= 20 and  # Minimum abstract length
                    paper.quality_score >= 1.0):  # Lower quality threshold
                    filtered_papers.append(paper)
        
        logger.info(f"Final filtered count: {len(filtered_papers)}")
        
        # Sort final results by combined score
        filtered_papers.sort(key=lambda p: (p.quality_score + p.relevance_score + (p.citations/100)), reverse=True)
        
        # Limit to maximum papers
        max_papers = self.config.get("search.max_papers", 25)
        return filtered_papers[:max_papers]

class ContentExtractor:
    """Enhanced content extraction with better insight generation"""
    
    def __init__(self):
        # Download required NLTK data
        try:
            nltk.data.find('tokenizers/punkt')
            nltk.data.find('corpora/stopwords')
        except LookupError:
            nltk.download('punkt', quiet=True)
            nltk.download('stopwords', quiet=True)
        
        self.stop_words = set(stopwords.words('english'))
    
    def extract_key_findings(self, paper: ResearchPaper) -> List[str]:
        """Enhanced key findings extraction"""
        if not paper.abstract:
            return []
        
        sentences = sent_tokenize(paper.abstract)
        key_sentences = []
        
        # Enhanced key indicators with weights
        indicators = {
            'high': ['found', 'discovered', 'revealed', 'demonstrated', 'concluded', 'results show'],
            'medium': ['showed', 'indicated', 'suggested', 'evidence', 'significant', 'correlation'],
            'low': ['relationship', 'effect', 'impact', 'influence', 'associated', 'related']
        }
        
        scored_sentences = []
        for sentence in sentences:
            sentence_lower = sentence.lower()
            score = 0
            
            # Score based on indicators
            for weight, words in [('high', indicators['high']), ('medium', indicators['medium']), ('low', indicators['low'])]:
                for word in words:
                    if word in sentence_lower:
                        if weight == 'high':
                            score += 3
                        elif weight == 'medium':
                            score += 2
                        else:
                            score += 1
                        break
            
            # Bonus for numerical data
            if re.search(r'\d+\.?\d*%|\d+\.?\d*\s*(fold|times|percent)', sentence_lower):
                score += 2
            
            # Bonus for statistical terms
            if any(term in sentence_lower for term in ['p <', 'significant', 'correlation', 'regression']):
                score += 1
            
            if score > 0:
                scored_sentences.append((sentence.strip(), score))
        
        # Sort by score and take top sentences
        scored_sentences.sort(key=lambda x: x[1], reverse=True)
        key_sentences = [sent[0] for sent in scored_sentences[:3]]
        
        # If no key sentences found, take informative sentences from end
        if not key_sentences and len(sentences) >= 2:
            key_sentences = sentences[-2:]
        elif not key_sentences and sentences:
            key_sentences = [sentences[-1]]
        
        return key_sentences
    
    def build_knowledge_context(self, papers: List[ResearchPaper]) -> Dict[str, Any]:
        """Enhanced knowledge context building"""
        if not papers:
            return {
                "total_papers": 0,
                "key_findings": [],
                "error": "No papers available for context building"
            }
        
        context = {
            "total_papers": len(papers),
            "key_findings": [],
            "common_themes": self._extract_themes(papers),
            "methodologies": self._extract_methodologies(papers),
            "recent_trends": self._identify_trends(papers),
            "citation_summary": {
                "total_citations": sum(p.citations for p in papers),
                "avg_citations": statistics.mean([p.citations for p in papers]),
                "median_citations": statistics.median([p.citations for p in papers]),
                "most_cited": max(papers, key=lambda p: p.citations),
                "citation_distribution": self._analyze_citations(papers)
            },
            "temporal_analysis": {
                "year_range": f"{min(p.year for p in papers)}-{max(p.year for p in papers)}",
                "recent_papers": len([p for p in papers if datetime.now().year - p.year <= 3]),
                "by_decade": self._analyze_by_decade(papers)
            },
            "venues": self._analyze_venues(papers),
            "quality_metrics": {
                "avg_quality_score": statistics.mean([p.quality_score for p in papers]),
                "avg_relevance_score": statistics.mean([p.relevance_score for p in papers]),
                "high_quality_papers": len([p for p in papers if p.quality_score >= 4.0])
            },
            "top_authors": self._identify_top_authors(papers)
        }
        
        # Extract and categorize findings
        for paper in papers:
            findings = self.extract_key_findings(paper)
            paper.key_findings = findings
            
            for finding in findings:
                author_name = paper.authors[0].split()[-1] if paper.authors else 'Unknown'
                context["key_findings"].append({
                    "text": finding,
                    "author": author_name,
                    "year": paper.year,
                    "citations": paper.citations,
                    "source": paper.source
                })
        
        return context
    
    def _extract_themes(self, papers: List[ResearchPaper]) -> List[str]:
        """Extract common themes from paper titles and abstracts"""
        all_text = " ".join([p.title + " " + p.abstract for p in papers if p.title and p.abstract])
        
        # Simple theme extraction (could be enhanced with more sophisticated NLP)
        theme_keywords = [
            "machine learning", "artificial intelligence", "deep learning", "neural networks",
            "climate change", "sustainability", "renewable energy", "carbon",
            "healthcare", "medical", "treatment", "therapy", "diagnosis",
            "education", "learning", "teaching", "student", "performance",
            "social", "economic", "policy", "government", "public",
            "data", "analysis", "model", "algorithm", "framework"
        ]
        
        found_themes = []
        all_text_lower = all_text.lower()
        
        for theme in theme_keywords:
            if theme in all_text_lower:
                count = all_text_lower.count(theme)
                if count >= 2:  # Appears in at least 2 papers
                    found_themes.append(f"{theme} ({count} occurrences)")
        
        return found_themes[:10]  # Top 10 themes
    
    def _extract_methodologies(self, papers: List[ResearchPaper]) -> List[str]:
        """Extract common methodologies mentioned in papers"""
        methodology_keywords = [
            "systematic review", "meta-analysis", "randomized controlled trial", "survey",
            "interview", "case study", "experimental", "longitudinal", "cross-sectional",
            "qualitative", "quantitative", "mixed methods", "regression analysis",
            "statistical analysis", "content analysis"
        ]
        
        all_text = " ".join([p.abstract for p in papers if p.abstract]).lower()
        found_methods = []
        
        for method in methodology_keywords:
            if method in all_text:
                found_methods.append(method)
        
        return found_methods[:8]  # Top 8 methodologies
    
    def _identify_trends(self, papers: List[ResearchPaper]) -> List[str]:
        """Identify temporal trends in research"""
        current_year = datetime.now().year
        recent_papers = [p for p in papers if current_year - p.year <= 3]
        older_papers = [p for p in papers if current_year - p.year > 3]
        
        trends = []
        
        if recent_papers:
            recent_text = " ".join([p.title + " " + p.abstract for p in recent_papers]).lower()
            older_text = " ".join([p.title + " " + p.abstract for p in older_papers]).lower()
            
            emerging_terms = ["AI", "machine learning", "deep learning", "blockchain", "IoT", 
                             "sustainability", "climate", "digital transformation", "remote"]
            
            for term in emerging_terms:
                recent_count = recent_text.count(term.lower())
                older_count = older_text.count(term.lower()) if older_text else 0
                
                if recent_count > older_count * 1.5:  # 50% increase in recent years
                    trends.append(f"Increasing focus on {term}")
        
        return trends[:5]
    
    def _analyze_citations(self, papers: List[ResearchPaper]) -> Dict[str, int]:
        """Analyze citation distribution"""
        citations = [p.citations for p in papers]
        return {
            "low_cited": len([c for c in citations if c < 10]),
            "medium_cited": len([c for c in citations if 10 <= c < 100]),
            "high_cited": len([c for c in citations if c >= 100])
        }
    
    def _analyze_by_decade(self, papers: List[ResearchPaper]) -> Dict[str, int]:
        """Analyze papers by decade"""
        decades = {}
        for paper in papers:
            decade = (paper.year // 10) * 10
            decade_key = f"{decade}s"
            decades[decade_key] = decades.get(decade_key, 0) + 1
        return decades
    
    def _analyze_venues(self, papers: List[ResearchPaper]) -> Dict[str, Any]:
        """Analyze publication venues"""
        venues = {}
        for paper in papers:
            if paper.venue:
                venues[paper.venue] = venues.get(paper.venue, 0) + 1
        
        return {
            "total_venues": len(venues),
            "top_venues": sorted(venues.items(), key=lambda x: x[1], reverse=True)[:5],
            "venue_diversity": len(venues) / len(papers) if papers else 0
        }
    
    def _identify_top_authors(self, papers: List[ResearchPaper]) -> List[Dict[str, Any]]:
        """Identify most frequent authors"""
        author_counts = {}
        author_citations = {}
        
        for paper in papers:
            for author in paper.authors:
                if author:
                    author_counts[author] = author_counts.get(author, 0) + 1
                    author_citations[author] = author_citations.get(author, 0) + paper.citations
        
        top_authors = []
        for author, count in sorted(author_counts.items(), key=lambda x: x[1], reverse=True)[:10]:
            top_authors.append({
                "name": author,
                "paper_count": count,
                "total_citations": author_citations[author]
            })
        
        return top_authors

class ArticleGenerator:
    """Enhanced article generator with retry logic and better prompts"""
    
    def __init__(self, config: Config):
        self.config = config
        openai.api_key = config.get("apis.openai_api_key")
        self.model = config.get("generation.model", "gpt-4")
        self.fallback_model = config.get("generation.fallback_model", "gpt-3.5-turbo")
        self.retry_attempts = config.get("generation.retry_attempts", 3)
    
    def generate_section(self, section_type: str, context: Dict[str, Any], 
                        refined_topic: Dict[str, str], papers: List[ResearchPaper] = None) -> ArticleSection:
        """Enhanced section generation with retry logic and better error handling"""
        
        if not context.get("total_papers", 0):
            logger.warning(f"No papers available for {section_type} generation")
            return self._create_fallback_section(section_type, refined_topic)
        
        prompts = {
            "title": self._get_title_prompt(),
            "abstract": self._get_abstract_prompt(),
            "introduction": self._get_introduction_prompt(),
            "literature_review": self._get_literature_review_prompt(),
            "method": self._get_method_prompt(),
            "results": self._get_results_prompt(),
            "conclusion": self._get_conclusion_prompt()
        }
        
        if section_type not in prompts:
            raise ValueError(f"Unknown section type: {section_type}")
        
        # Enhanced context formatting
        formatted_context = self._format_context(context, papers)
        
        prompt = prompts[section_type].format(
            topic=refined_topic["title"],
            research_question=refined_topic["research_question"],
            context=formatted_context,
            target_words=self.config.get(f"generation.target_word_counts.{section_type}", 500)
        )
        
        # Try generation with retries
        for attempt in range(self.retry_attempts):
            try:
                from openai import OpenAI
                current_model = self.model if attempt < 2 else self.fallback_model
                openai_api_key = self.config.get("apis.openai_api_key")
                client = OpenAI(api_key=openai_api_key)

                response = client.chat.completions.create(
                model="gpt-5-mini", #current_model,
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are an expert academic writer specializing in research articles. "
                            "Write in formal academic style with proper citations. "
                            "Focus on clarity, coherence, and academic rigor."
                        )
                    },
                    {"role": "user", "content": prompt}
                ],
                temperature=self.config.get("generation.temperature", 1.0),
                max_completion_tokens=self.config.get("generation.max_completion_tokens", 3500)
                )

                content = response.choices[0].message.content.strip()
                
                # Validate generated content
                if self._validate_content(content, section_type):
                    return ArticleSection(
                        title=section_type.replace("_", " ").title(),
                        content=content
                    )
                else:
                    logger.warning(f"Generated content for {section_type} failed validation, retrying...")
                    
            except Exception as e:
                import sys
                lineno = sys.exc_info()[2].tb_lineno
                logger.error(f"Error generating {section_type} (attempt {attempt + 1}) at line {lineno}: {e}")
                if attempt == self.retry_attempts - 1:
                    logger.error(f"All attempts failed for {section_type}")
                    return self._create_fallback_section(section_type, refined_topic)
                time.sleep(2 ** attempt)  # Exponential backoff
        
        return self._create_fallback_section(section_type, refined_topic)
    
    def _validate_content(self, content: str, section_type: str) -> bool:
        """Validate generated content quality"""
        if not content or len(content.strip()) < 50:
            return False
        
        word_count = len(content.split())
        min_words = self.config.get("quality.min_section_words", 100)
        max_words = self.config.get("quality.max_section_words", 2500)
        
        if word_count < min_words or word_count > max_words:
            logger.warning(f"{section_type} word count ({word_count}) outside acceptable range")
            return False
        
        # Check for placeholder text
        placeholders = ["[insert", "TODO", "placeholder", "xxx", "fill in"]
        content_lower = content.lower()
        if any(placeholder in content_lower for placeholder in placeholders):
            return False
        
        return True
    
    def _create_fallback_section(self, section_type: str, refined_topic: Dict[str, str]) -> ArticleSection:
        """Create fallback content when generation fails"""
        fallback_content = {
            "abstract": f"This study examines {refined_topic['original_topic']}. Due to technical limitations, detailed analysis was not generated. This research area requires further investigation to provide comprehensive insights.",
            "introduction": f"The topic of {refined_topic['original_topic']} represents an important area of research. This introduction provides a foundation for understanding the key concepts and significance of this field of study.",
            "literature_review": f"A comprehensive review of literature on {refined_topic['original_topic']} reveals various perspectives and findings. Further detailed analysis of relevant studies would enhance understanding of current research trends.",
            "method": f"The methodology for investigating {refined_topic['original_topic']} requires careful consideration of appropriate research designs, data collection methods, and analysis techniques suitable for this field of study.",
            "results": f"Results related to {refined_topic['original_topic']} would typically present findings from systematic investigation. Detailed analysis and interpretation of data would be presented in this section.",
            "conclusion": f"In conclusion, {refined_topic['original_topic']} presents significant opportunities for research and application. Future studies should focus on expanding understanding and addressing current limitations."
        }
        
        content = fallback_content.get(section_type, f"Content for {section_type} section on {refined_topic['original_topic']} requires further development.")
        
        return ArticleSection(
            title=section_type.replace("_", " ").title(),
            content=content
        )
    
    def _format_context(self, context: Dict[str, Any], papers: List[ResearchPaper] = None) -> str:
        """Enhanced context formatting for LLM prompt"""
        if context.get("error"):
            return f"Limited research context available. {context.get('error')}"
        
        formatted = f"""
Research Context Summary:
- Total Papers Analyzed: {context['total_papers']}
- Citation Statistics: 
  * Total Citations: {context['citation_summary']['total_citations']}
  * Average Citations: {context['citation_summary']['avg_citations']:.1f}
  * Median Citations: {context['citation_summary']['median_citations']:.1f}
- Time Span: {context['temporal_analysis']['year_range']}
- Recent Research: {context['temporal_analysis']['recent_papers']} papers from last 3 years
- Quality Metrics:
  * Average Quality Score: {context['quality_metrics']['avg_quality_score']:.2f}
  * High Quality Papers: {context['quality_metrics']['high_quality_papers']}

"""
        
        # Add key findings
        if context.get('key_findings'):
            formatted += "Key Research Findings:\n"
            for i, finding in enumerate(context['key_findings'][:15], 1):  # Top 15 findings
                if isinstance(finding, dict):
                    formatted += f"{i}. {finding['text']} ({finding['author']}, {finding['year']})\n"
                else:
                    formatted += f"{i}. {finding}\n"
            formatted += "\n"
        
        # Add common themes
        if context.get('common_themes'):
            formatted += f"Common Themes: {', '.join(context['common_themes'][:5])}\n\n"
        
        # Add methodologies
        if context.get('methodologies'):
            formatted += f"Common Methodologies: {', '.join(context['methodologies'][:5])}\n\n"
        
        # Add recent trends
        if context.get('recent_trends'):
            formatted += f"Recent Trends: {', '.join(context['recent_trends'])}\n\n"
        
        return formatted
    
    def _get_title_prompt(self) -> str:
        return """
Generate a compelling, academic title for a research article on "{topic}".
The title should be:
- Precise and descriptive (10-15 words)
- Following academic conventions
- Engaging but professional
- Reflecting current research trends

Research Context: {context}

Generate only the title, no additional text.
"""
    
    def _get_abstract_prompt(self) -> str:
        return """
Write a comprehensive abstract for a research article titled "{topic}".
Target length: approximately {target_words} words.

Structure the abstract with these elements:
1. Background/Context (2-3 sentences)
2. Research objective/question: {research_question}
3. Methodology (brief, 1-2 sentences)
4. Key findings (2-3 sentences)
5. Implications/Conclusions (1-2 sentences)

Research Context: {context}

Write in formal academic style following APA7 format. Do not include citations in the abstract.
Ensure the abstract is self-contained and provides a complete overview of the research.
"""
    
    def _get_introduction_prompt(self) -> str:
        return """
Write a comprehensive introduction for a research article titled "{topic}".
Target length: approximately {target_words} words.

Structure the introduction as follows:
1. Broad context and background of the research area
2. Narrowing focus to specific research problem
3. Literature gap or research need identification
4. Research question and objectives: {research_question}
5. Significance and potential contributions
6. Brief overview of article structure

Research Context: {context}

Requirements:
- Use formal academic writing style
- Include relevant in-text citations in APA7 format (Author, Year)
- Build a logical argument leading to the research question
- Demonstrate knowledge of current literature
- Maintain scholarly tone throughout
"""
    
    def _get_literature_review_prompt(self) -> str:
        return """
Write a comprehensive literature review for a research article on "{topic}".
Target length: approximately {target_words} words.

Organize the literature review with:
1. Thematic organization of existing research
2. Synthesis of key findings and methodologies
3. Critical analysis of strengths and limitations
4. Identification of research gaps
5. Theoretical framework development
6. Connection to current research question: {research_question}

Research Context: {context}

Requirements:
- Use extensive citations in APA7 format throughout
- Organize into clear thematic subsections with appropriate headings
- Provide critical analysis, not just summary
- Synthesize rather than just list studies
- Build toward justification for current research
- Use transitional sentences to connect ideas
"""
    
    def _get_method_prompt(self) -> str:
        return """
Write a detailed methodology section for a research article on "{topic}".
Target length: approximately {target_words} words.

Include these subsections:
1. Research Design and Approach
2. Participants/Sample (if applicable)
3. Data Collection Methods and Procedures
4. Instruments/Materials
5. Data Analysis Plan
6. Validity and Reliability Considerations
7. Ethical Considerations
8. Limitations

Research Question Focus: {research_question}
Research Context: {context}

Requirements:
- Be specific about procedures and justify methodological choices
- Use appropriate research terminology
- Provide sufficient detail for replication
- Address potential threats to validity
- Follow standard academic format
"""
    
    def _get_results_prompt(self) -> str:
        return """
Write a comprehensive results section for a research article on "{topic}".
Target length: approximately {target_words} words.

Structure the results section:
1. Overview of analysis approach
2. Descriptive statistics/preliminary findings
3. Main findings organized by research questions
4. Statistical results with appropriate reporting
5. Additional relevant findings
6. Summary of key results

Research Question: {research_question}
Research Context: {context}

Requirements:
- Present findings objectively without interpretation
- Use appropriate statistical language and reporting standards
- Reference tables and figures (describe what they would contain)
- Organize results logically
- Include effect sizes and confidence intervals where appropriate
- Maintain clear, concise academic writing
"""
    
    def _get_conclusion_prompt(self) -> str:
        return """
Write a comprehensive conclusion section for a research article on "{topic}".
Target length: approximately {target_words} words.

Structure the conclusion:
1. Summary of key findings
2. Direct response to research question: {research_question}
3. Interpretation of results in broader context
4. Theoretical and practical implications
5. Limitations and their implications
6. Recommendations for future research
7. Final concluding thoughts on significance

Research Context: {context}

Requirements:
- Synthesize findings rather than just restating results
- Connect findings to existing literature
- Discuss broader implications for theory and practice
- Acknowledge study limitations honestly
- Provide specific, actionable recommendations for future research
- End with strong concluding statement about contribution to field
"""

class CitationManager:
    """Enhanced citation management with better formatting"""
    
    def __init__(self):
        self.references = {}
        self.citation_count = 0
    
    def add_reference(self, paper: ResearchPaper) -> str:
        """Add a reference and return citation key"""
        # Generate citation key
        if paper.authors:
            first_author_last = paper.authors[0].split()[-1] if paper.authors[0] else "Unknown"
        else:
            first_author_last = "Unknown"
        
        key = f"{first_author_last}{paper.year}"
        
        # Handle duplicates
        original_key = key
        counter = 1
        while key in self.references:
            key = f"{original_key}_{counter}"
            counter += 1
        
        self.references[key] = paper
        return key
    
    def generate_bibliography(self) -> str:
        """Generate enhanced APA7 formatted bibliography"""
        if not self.references:
            return "No references available."
        
        bibliography = []
        
        for key, paper in sorted(self.references.items(), key=lambda x: x[1].authors[0].split()[-1] if x[1].authors else ""):
            try:
                citation = self._format_apa_citation(paper)
                bibliography.append(citation)
            except Exception as e:
                import sys
                logger.warning(f"Error formatting citation for {key}: {e} (line {sys.exc_info()[2].tb_lineno})")
                # Fallback simple format
                author_str = ", ".join(paper.authors[:3]) if paper.authors else "Unknown Author"
                if len(paper.authors) > 3:
                    author_str += " et al."
                citation = f"{author_str} ({paper.year}). {paper.title}."
                bibliography.append(citation)
        
        return "\n\n".join(bibliography)
    
    def _format_apa_citation(self, paper: ResearchPaper) -> str:
        """Format a single citation in APA7 style"""
        # Format authors
        if paper.authors:
            if len(paper.authors) == 1:
                author_str = paper.authors[0]
            elif len(paper.authors) == 2:
                author_str = f"{paper.authors[0]} & {paper.authors[1]}"
            elif len(paper.authors) <= 20:
                author_str = ", ".join(paper.authors[:-1]) + f", & {paper.authors[-1]}"
            else:
                author_str = ", ".join(paper.authors[:19]) + f", ... {paper.authors[-1]}"
        else:
            author_str = "Unknown Author"
        
        # Start citation
        citation = f"{author_str} ({paper.year}). {paper.title}"
        
        # Add period if title doesn't end with punctuation
        if not citation.endswith('.'):
            citation += "."
        
        # Add venue information
        if paper.venue:
            if paper.source == "arXiv":
                citation += f" *arXiv preprint*. "
            else:
                citation += f" *{paper.venue}*. "
        
        # Add DOI or URL
        if paper.doi:
            citation += f"https://doi.org/{paper.doi}"
        elif paper.url and paper.url.startswith('http'):
            citation += f"Retrieved from {paper.url}"
        
        return citation

class DocumentFormatter:
    """Enhanced document formatter with better styling"""
    
    def __init__(self, config: Config):
        self.config = config
        self.output_dir = Path(config.get("output.output_dir", "outputs"))
        self.output_dir.mkdir(exist_ok=True)
        
    def _setup_document_styles(self, doc):
        
        """Setup document styles for better formatting"""
        try:
            # This is a basic setup - could be enhanced with more sophisticated styling
            pass
        except Exception as e:
            import sys
            logger.warning(f"Could not setup document styles: {e} (line {sys.exc_info()[2].tb_lineno})")
    
    def create_docx(self, title: str, sections: List[ArticleSection], 
                   bibliography: str, keywords: List[str]) -> str:
        """Create enhanced APA7 formatted Word document"""
        doc = Document()
        
        # Configure document styles
        self._setup_document_styles(doc)
        
        # Title page
        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.runs[0].bold = True
        
        doc.add_paragraph()  # Spacing
        
        # Author information (placeholder)
        author_para = doc.add_paragraph("Author Name\nInstitution Name\nEmail: author@institution.edu")
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Page break
        doc.add_page_break()
        
        # Abstract
        abstract_section = next((s for s in sections if "abstract" in s.title.lower()), None)
        if abstract_section:
            abstract_heading = doc.add_paragraph("Abstract")
            abstract_heading.runs[0].bold = True
            abstract_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph(abstract_section.content)
            doc.add_paragraph()
        
        # Keywords
        if keywords:
            keywords_para = doc.add_paragraph(f"Keywords: {', '.join(keywords)}")
            keywords_para.runs[0].italic = True
            doc.add_paragraph()
        
        # Page break before main content
        doc.add_page_break()
        
        # Main sections
        for section in sections:
            if "abstract" not in section.title.lower():
                heading = doc.add_paragraph(section.title)
                heading.runs[0].bold = True
                doc.add_paragraph(section.content)
                doc.add_paragraph()
        
        # References
        if bibliography:
            ref_heading = doc.add_paragraph("References")
            ref_heading.runs[0].bold = True
            ref_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(bibliography)
        
        # Save document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = re.sub(r'[^\w\s-]', '', title)[:30].replace(' ', '_')
        filename = f"{safe_title}_{timestamp}.docx"
        filepath = self.output_dir / filename
        
        try:
            doc.save(str(filepath))
            logger.info(f"Word document saved: {filepath}")
            return str(filepath)
        except Exception as e:
            import sys
            logger.error(f"Failed to save Word document: {e} (line {sys.exc_info()[2].tb_lineno})")
            raise
            
    def create_markdown(self, title: str, sections: List[ArticleSection], 
                       bibliography: str, keywords: List[str], 
                       context: Dict[str, Any] = None) -> str:
        """Create enhanced Markdown version with metadata"""
        content = [f"# {title}\n"]
        
        # Add metadata
        content.append("---")
        content.append(f"title: \"{title}\"")
        content.append(f"generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        if keywords:
            content.append(f"keywords: [{', '.join(f'\"{k}\"' for k in keywords)}]")
        if context:
            content.append(f"papers_analyzed: {context.get('total_papers', 0)}")
        content.append("---\n")
        
        # Keywords
        if keywords:
            content.append(f"**Keywords:** {', '.join(keywords)}\n")
        
        # Table of Contents
        content.append("## Table of Contents\n")
        for section in sections:
            section_link = section.title.lower().replace(' ', '-').replace('_', '-')
            content.append(f"- [{section.title}](#{section_link})")
        content.append("- [References](#references)\n")
        
        # Sections
        for section in sections:
            content.append(f"## {section.title}\n")
            content.append(f"{section.content}\n")
        
        # References
        if bibliography:
            content.append("## References\n")
            content.append(f"{bibliography}\n")
        
        # Statistics footer
        if context:
            content.append("---\n")
            content.append("### Article Statistics\n")
            total_words = sum(s.word_count for s in sections)
            content.append(f"- **Total Words:** {total_words:,}")
            content.append(f"- **Papers Analyzed:** {context.get('total_papers', 0)}")
            content.append(f"- **Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Save markdown
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = re.sub(r'[^\w\s-]', '', title)[:30].replace(' ', '_')
        filename = f"{safe_title}_{timestamp}.md"
        filepath = self.output_dir / filename
        
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write('\n'.join(content))
            logger.info(f"Markdown saved: {filepath}")
            return str(filepath)
        except Exception as e:
            import sys
            logger.error(f"Error saving Markdown: {e} (line {sys.exc_info()[2].tb_lineno})")
            # Try with simpler filename
            filename = f"research_article_{timestamp}.md"
            filepath = self.output_dir / filename
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write('\n'.join(content))
            return str(filepath)

class ResearchArticleGenerator:
    """Enhanced main orchestrator class with better error handling"""
    
    def __init__(self, config_path: str = "config.yaml"):
        self.config = Config(config_path)
        self.searcher = PaperSearcher(self.config)
        self.extractor = ContentExtractor()
        self.generator = ArticleGenerator(self.config)
        self.citation_manager = CitationManager()
        self.formatter = DocumentFormatter(self.config)
        
        # Validate setup
        self._validate_setup()
    
    def _validate_setup(self):
        """Validate that the generator is properly set up"""
        if not self.config.get("apis.openai_api_key"):
            logger.error("OpenAI API key is required for article generation")
            raise ValueError("OpenAI API key not found. Please set OPENAI_API_KEY environment variable.")
    
    def generate_article(self, topic: str) -> Dict[str, Any]:
        """Enhanced main method to generate complete research article"""
        start_time = time.time()
        logger.info(f"Starting research article generation for topic: '{topic}'")
        try:
            # Step 1: Refine topic
            logger.info("Step 1: Refining research topic...")
            refined_topic = TopicRefiner.refine_topic(topic)
            logger.info(f"Refined title: {refined_topic['title']}")
            logger.info(f"Research question: {refined_topic['research_question']}")
            # Step 2: Search for papers
            logger.info("Step 2: Searching for relevant papers...")
            search_query = " ".join(refined_topic['search_terms'])
            papers = self.searcher.search_all_sources(search_query)
            # Enhanced fallback handling
            if not papers:
                logger.warning("No papers found with primary search terms")
                # Try with broader terms
                broader_query = " ".join(refined_topic['search_terms'][:3])
                logger.info(f"Attempting broader search: '{broader_query}'")
                papers = self.searcher.search_all_sources(broader_query)
            if not papers:
                logger.warning("Still no papers found. Generating article with limited context...")
                return self._generate_limited_article(refined_topic)
            logger.info(f"Using {len(papers)} papers for article generation")
            # Add papers to citation manager
            for paper in papers:
                self.citation_manager.add_reference(paper)
            # Step 3: Extract knowledge context
            logger.info("Step 3: Extracting knowledge context...")
            context = self.extractor.build_knowledge_context(papers)
            # Step 4: Generate article sections
            logger.info("Step 4: Generating article sections...")
            sections = []
            section_types = ["abstract", "introduction", "literature_review", "method", "results", "conclusion"]
            for section_type in tqdm(section_types, desc="Generating sections"):
                logger.info(f"Generating {section_type}...")
                try:
                    section = self.generator.generate_section(section_type, context, refined_topic, papers)
                    sections.append(section)
                    logger.info(f"Successfully generated {section_type} ({section.word_count} words)")
                except Exception as e:
                    import sys
                    logger.error(f"Failed to generate {section_type}: {e} (line {sys.exc_info()[2].tb_lineno})")
                    # Create fallback section
                    fallback_section = self.generator._create_fallback_section(section_type, refined_topic)
                    sections.append(fallback_section)
            # Step 5: Generate bibliography
            logger.info("Step 5: Generating bibliography...")
            bibliography = self.citation_manager.generate_bibliography()
            # Step 6: Generate keywords
            logger.info("Step 6: Generating keywords...")
            keywords = self._generate_keywords(refined_topic, context)
            # Step 7: Format and save documents
            logger.info("Step 7: Creating output documents...")
            output_files = {}
            # Create Word document
            if "docx" in self.config.get("output.format", ["docx"]):
                try:
                    docx_path = self.formatter.create_docx(
                        refined_topic["title"], sections, bibliography, keywords
                    )
                    output_files["docx"] = docx_path
                except Exception as e:
                    logger.error(f"Failed to create Word document: {e}")
            # Create Markdown document
            if "markdown" in self.config.get("output.format", ["docx"]):
                try:
                    md_path = self.formatter.create_markdown(
                        refined_topic["title"], sections, bibliography, keywords, context
                    )
                    output_files["markdown"] = md_path
                except Exception as e:
                    logger.error(f"Failed to create Markdown document: {e}")
            # Create summary report
            if self.config.get("output.include_summary", True):
                try:
                    summary_path = self._create_summary_report(refined_topic, context, sections, papers)
                    output_files["summary"] = summary_path
                except Exception as e:
                    logger.error(f"Failed to create summary report: {e}")
            # Calculate generation time
            generation_time = time.time() - start_time
            # Compile results
            result = {
                "status": "success",
                "title": refined_topic["title"],
                "files": output_files,
                "stats": {
                    "papers_analyzed": len(papers),
                    "total_words": sum(s.word_count for s in sections),
                    "sections_generated": len(sections),
                    "references": len(self.citation_manager.references),
                    "generation_time_minutes": round(generation_time / 60, 2),
                    "search_stats": self.searcher.search_stats,
                    "quality_metrics": self._calculate_quality_metrics(sections, context)
                },
                "warnings": self._collect_warnings()
            }
            logger.info(f"Article generation completed successfully in {generation_time/60:.2f} minutes!")
            return result
        except Exception as e:
            logger.error(f"Article generation failed: {e}")
            return {
                "status": "error",
                "error": str(e),
                "title": refined_topic.get("title", topic) if 'refined_topic' in locals() else topic,
                "generation_time_minutes": round((time.time() - start_time) / 60, 2)
            }
    
    def _generate_limited_article(self, refined_topic: Dict[str, str]) -> Dict[str, Any]:
        """Generate article with limited context when no papers are found"""
        logger.warning("Generating article with limited research context")
        
        # Create minimal context
        context = {
            "total_papers": 0,
            "key_findings": [],
            "error": "No research papers found for analysis"
        }
        
        # Generate basic sections
        sections = []
        section_types = ["abstract", "introduction", "literature_review", "conclusion"]
        
        for section_type in section_types:
            section = self.generator._create_fallback_section(section_type, refined_topic)
            sections.append(section)
        
        # Generate minimal bibliography
        bibliography = "No references available due to limited paper search results."
        
        # Generate basic keywords
        keywords = refined_topic['search_terms'][:5]
        
        # Create output files
        output_files = {}
        try:
            if "markdown" in self.config.get("output.format", ["markdown"]):
                md_path = self.formatter.create_markdown(
                    refined_topic["title"], sections, bibliography, keywords, context
                )
                output_files["markdown"] = md_path
        except Exception as e:
            import sys
            logger.error(f"Failed to create limited article: {e} (line {sys.exc_info()[2].tb_lineno})")
        
        return {
            "status": "limited_success",
            "title": refined_topic["title"],
            "files": output_files,
            "stats": {
                "papers_analyzed": 0,
                "total_words": sum(s.word_count for s in sections),
                "sections_generated": len(sections),
                "references": 0
            },
            "warnings": ["No research papers found", "Generated with limited context"]
        }
    
    def _calculate_quality_metrics(self, sections: List[ArticleSection], context: Dict[str, Any]) -> Dict[str, Any]:
        """Calculate quality metrics for the generated article"""
        total_words = sum(s.word_count for s in sections)
        avg_section_quality = statistics.mean([s.quality_score for s in sections]) if sections else 0
        
        # Calculate readability
        all_text = " ".join([s.content for s in sections])
        try:
            readability_score = flesch_reading_ease(all_text)
        except:
            readability_score = 0
        
        return {
            "total_words": total_words,
            "average_section_quality": round(avg_section_quality, 2),
            "readability_score": round(readability_score, 1),
            "research_foundation_strength": min(10, context.get('total_papers', 0) / 2),  # Scale of 0-10
            "section_completeness": len(sections) / 6 * 100  # Percentage of expected sections
        }
    
    def _collect_warnings(self) -> List[str]:
        """Collect any warnings that occurred during generation"""
        warnings = []
        
        # Check for search warnings
        if self.searcher.search_stats.get("after_filtering", 0) < 10:
            warnings.append("Low number of papers found - consider broader search terms")
        
        # Check for API key warnings
        if not self.config.get("apis.openai_api_key"):
            warnings.append("OpenAI API key not configured")
        
        return warnings
    
    def _generate_keywords(self, refined_topic: Dict[str, str], context: Dict[str, Any]) -> List[str]:
        """Enhanced keyword generation"""
        keywords = set()
        
        # From topic
        keywords.update(refined_topic['search_terms'])
        
        # From common themes
        if context.get('common_themes'):
            for theme in context['common_themes'][:3]:
                # Extract the theme name (before parentheses if present)
                theme_name = theme.split('(')[0].strip()
                keywords.add(theme_name)
        
        # From methodologies
        if context.get('methodologies'):
            keywords.update(context['methodologies'][:2])
        
        # Academic terms
        academic_terms = ['research', 'analysis', 'study', 'investigation', 'evaluation']
        keywords.update([term for term in academic_terms if term in refined_topic['original_topic'].lower()])
        
        # Clean and format keywords
        cleaned_keywords = []
        for keyword in keywords:
            if isinstance(keyword, str) and len(keyword.split()) <= 3 and len(keyword) > 2:
                cleaned_keywords.append(keyword.lower())
        
        # Remove duplicates and limit
        final_keywords = list(set(cleaned_keywords))[:7]
        
        return final_keywords
    
    def _create_summary_report(self, refined_topic: Dict[str, str], context: Dict[str, Any], 
                              sections: List[ArticleSection], papers: List[ResearchPaper]) -> str:
        """Create enhanced summary report"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"generation_report_{timestamp}.md"
        filepath = self.formatter.output_dir / filename
        
        report_content = f"""# Research Article Generation Report

## Generation Summary
- **Original Topic:** {refined_topic['original_topic']}
- **Refined Title:** {refined_topic['title']}
- **Research Question:** {refined_topic['research_question']}
- **Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
- **Generator Version:** Enhanced Article Generator v2.0

## Literature Analysis
- **Papers Found:** {context.get('total_papers', 0)}
- **Papers by Source:**
"""
        
        # Add source breakdown
        for source, count in self.searcher.search_stats.get('by_source', {}).items():
            report_content += f"  * {source}: {count} papers\n"
        
        if context.get('total_papers', 0) > 0:
            report_content += f"""
- **Citation Statistics:**
  * Total Citations: {context['citation_summary']['total_citations']:,}
  * Average Citations: {context['citation_summary']['avg_citations']:.1f}
  * Median Citations: {context['citation_summary']['median_citations']:.1f}
  * Most Cited Paper: {context['citation_summary']['most_cited'].citations} citations
- **Temporal Distribution:** {context['temporal_analysis']['year_range']}
- **Recent Research:** {context['temporal_analysis']['recent_papers']} papers from last 3 years
- **Quality Metrics:**
  * Average Quality Score: {context['quality_metrics']['avg_quality_score']:.2f}/5.0
  * High Quality Papers: {context['quality_metrics']['high_quality_papers']}
"""
        
        # Article statistics
        report_content += f"""
## Generated Article Statistics
"""
        
        total_words = 0
        for section in sections:
            report_content += f"- **{section.title}:** {section.word_count:,} words (Quality: {section.quality_score:.1f}/5.0)\n"
            total_words += section.word_count
        
        report_content += f"- **Total Words:** {total_words:,}\n"
        
        # Quality assessment
        quality_metrics = self._calculate_quality_metrics(sections, context)
        report_content += f"""
## Quality Assessment
- **Overall Readability:** {quality_metrics['readability_score']} (Flesch Reading Ease)
- **Research Foundation:** {quality_metrics['research_foundation_strength']:.1f}/10.0
- **Section Completeness:** {quality_metrics['section_completeness']:.1f}%
- **Average Section Quality:** {quality_metrics['average_section_quality']:.1f}/5.0

### Quality Interpretation
- **Readability Score:** 
  * 90-100: Very Easy
  * 80-89: Easy  
  * 70-79: Fairly Easy
  * 60-69: Standard
  * 50-59: Fairly Difficult
  * 30-49: Difficult
  * 0-29: Very Difficult

"""
        
        # Top referenced papers
        if papers:
            report_content += "## Top Referenced Papers\n"
            top_papers = sorted(papers, key=lambda p: (p.citations, p.quality_score), reverse=True)[:5]
            for i, paper in enumerate(top_papers, 1):
                authors_str = ", ".join(paper.authors[:2]) if paper.authors else "Unknown"
                if len(paper.authors) > 2:
                    authors_str += " et al."
                report_content += f"""
{i}. **{paper.title}**
   - Authors: {authors_str}
   - Year: {paper.year}
   - Citations: {paper.citations:,}
   - Source: {paper.source}
   - Quality Score: {paper.quality_score:.1f}/5.0
"""
        
        # Recommendations
        report_content += f"""
## Review Recommendations

### Critical Review Tasks
1. **Verify Citations:** Cross-reference all in-text citations with the reference list
2. **Methodology Validation:** Ensure the methodology section matches your actual research approach
3. **Results Accuracy:** Replace generated results with your actual findings and data
4. **Literature Coverage:** Add any missing key references in your field
5. **Claim Verification:** Validate all factual claims against original sources

### Enhancement Suggestions
1. **Add Specificity:** Include specific data, figures, tables, and statistical analyses
2. **Customize Context:** Adapt content to your specific research context and objectives  
3. **Technical Details:** Add technical specifications, parameters, and detailed procedures
4. **Visual Elements:** Create appropriate figures, charts, and tables to support findings
5. **Field-Specific Terms:** Incorporate specialized terminology relevant to your discipline

### Quality Assurance
1. **Plagiarism Check:** Run content through plagiarism detection software
2. **Grammar Review:** Perform thorough proofreading and editing
3. **Format Compliance:** Ensure adherence to target journal's formatting requirements
4. **Peer Review:** Have colleagues review for accuracy and clarity
5. **Citation Format:** Verify all citations follow APA7 guidelines correctly

## Common Themes Identified
"""
        
        if context.get('common_themes'):
            for theme in context['common_themes']:
                report_content += f"- {theme}\n"
        
        if context.get('methodologies'):
            report_content += f"""
## Research Methodologies Mentioned
"""
            for method in context['methodologies']:
                report_content += f"- {method}\n"
        
        # Warnings and limitations
        warnings = self._collect_warnings()
        if warnings:
            report_content += f"""
## Warnings and Limitations
"""
            for warning in warnings:
                report_content += f"- ⚠️ {warning}\n"
        
        report_content += f"""
## Files Generated
- Research Article (Word format): Available if configured
- Research Article (Markdown format): Available if configured
- Generation Report: {filename}

## Next Steps Checklist
- [ ] Review abstract for accuracy and completeness
- [ ] Verify introduction provides adequate context
- [ ] Validate literature review covers key sources
- [ ] Customize methodology to match actual research
- [ ] Replace results with actual findings
- [ ] Ensure conclusion addresses research questions
- [ ] Add figures, tables, and appendices as needed
- [ ] Proofread entire document
- [ ] Check citation format and reference list
- [ ] Review for journal-specific requirements
- [ ] Run plagiarism check
- [ ] Obtain peer feedback before submission

---
*Generated by Enhanced Research Article Generator v2.0*  
*Report generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*
"""
        
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(report_content)
            logger.info(f"Enhanced summary report saved: {filepath}")
            return str(filepath)
        except Exception as e:
            import sys
            logger.error(f"Failed to create summary report: {e} (line {sys.exc_info()[2].tb_lineno})")
            return ""

def main():
    """Enhanced CLI interface with better error handling and options"""
    parser = argparse.ArgumentParser(
        description="Enhanced Automated Research Article Generator",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python articlegen.py "machine learning in healthcare"
  python articlegen.py "climate change impact" --config custom_config.yaml
  python articlegen.py "social media effects" --output ./my_articles --verbose
        """
    )
    
    parser.add_argument("topic", help="Research topic to generate article for")
    parser.add_argument("--config", default="config.yaml", help="Configuration file path")
    parser.add_argument("--output", default="outputs", help="Output directory")
    parser.add_argument("--verbose", "-v", action="store_true", help="Verbose logging")
    parser.add_argument("--format", choices=["docx", "markdown", "both"], default="markdown", #"both",
                       help="Output format (default: both)")
    parser.add_argument("--max-papers", type=int, help="Maximum number of papers to analyze")
    parser.add_argument("--no-summary", action="store_true", help="Skip generation report")
    
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
        logger.debug("Verbose logging enabled")
    
    logger.info("Starting Enhanced Research Article Generator")
    logger.info(f"Topic: {args.topic}")
    
    try:
        # Initialize generator
        generator = ResearchArticleGenerator(args.config)
        
        # Update configuration based on CLI args
        if args.output != "outputs":
            generator.formatter.output_dir = Path(args.output)
            generator.formatter.output_dir.mkdir(exist_ok=True)
        
        if args.format != "both":
            generator.config.config["output"]["format"] = [args.format]
        
        if args.max_papers:
            generator.config.config["search"]["max_papers"] = args.max_papers
        
        if args.no_summary:
            generator.config.config["output"]["include_summary"] = False
        
        # Generate article
        result = generator.generate_article(args.topic)
        
        # Handle results
        if result["status"] == "error":
            print(f"❌ Error: {result['error']}")
            print(f"⏱️  Generation time: {result.get('generation_time_minutes', 0):.2f} minutes")
            sys.exit(1)
        
        elif result["status"] == "limited_success":
            print("\n⚠️  Article Generation Completed with Limitations!")
            print(f"📝 Title: {result['title']}")
            print(f"⚠️  Warning: Generated with limited research context")
            print(f"📊 Statistics:")
            print(f"   - Papers analyzed: {result['stats']['papers_analyzed']}")
            print(f"   - Total words: {result['stats']['total_words']:,}")
            print(f"   - Sections generated: {result['stats']['sections_generated']}")
            
        else:
            # Successful generation
            print("\n🎉 Article Generation Completed Successfully!")
            print(f"📝 Title: {result['title']}")
            print(f"📊 Statistics:")
            print(f"   - Papers analyzed: {result['stats']['papers_analyzed']:,}")
            print(f"   - Total words: {result['stats']['total_words']:,}")
            print(f"   - Sections generated: {result['stats']['sections_generated']}")
            print(f"   - References: {result['stats']['references']}")
            print(f"   - Generation time: {result['stats']['generation_time_minutes']:.2f} minutes")
            
            # Quality metrics
            if 'quality_metrics' in result['stats']:
                qm = result['stats']['quality_metrics']
                print(f"📈 Quality Metrics:")
                print(f"   - Research foundation: {qm['research_foundation_strength']:.1f}/10.0")
                print(f"   - Readability score: {qm['readability_score']:.1f}")
                print(f"   - Section completeness: {qm['section_completeness']:.1f}%")
        
        # Show generated files
        if result.get("files"):
            print(f"\n📁 Generated Files:")
            for format_type, filepath in result['files'].items():
                print(f"   - {format_type.upper()}: {filepath}")
        
        # Show warnings
        if result.get("warnings"):
            print(f"\n⚠️  Warnings:")
            for warning in result['warnings']:
                print(f"   - {warning}")
        
        print(f"\n✅ Your research article is ready for review and customization!")
        if result.get("files", {}).get("summary"):
            print(f"💡 Please review the generation report for detailed recommendations.")
        
        # Search statistics
        if result.get('stats', {}).get('search_stats'):
            search_stats = result['stats']['search_stats']
            print(f"\n🔍 Search Statistics:")
            print(f"   - Total papers found: {search_stats.get('total_found', 0)}")
            print(f"   - After filtering: {search_stats.get('after_filtering', 0)}")
            if search_stats.get('by_source'):
                print(f"   - By source:")
                for source, count in search_stats['by_source'].items():
                    print(f"     * {source}: {count}")
        
    except KeyboardInterrupt:
        print("\n⚠️ Generation interrupted by user")
        sys.exit(1)
    except ValueError as e:
        print(f"❌ Configuration Error: {e}")
        print("💡 Tip: Make sure your OpenAI API key is set in the OPENAI_API_KEY environment variable")
        sys.exit(1)
    except Exception as e:
        logger.exception("Unexpected error during generation")
        print(f"❌ Unexpected error: {e}")
        print("💡 Check the logs for more details")
        sys.exit(1)

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        import sys
        logger.error(f"Error saving Word document: {e} (line {sys.exc_info()[2].tb_lineno})")
        # Try with simpler filename
        # filename = f"research_article_{timestamp}.docx"
        # filepath = self.output_dir / filename
        # doc.save(str(filepath))
        # return str(filepath)       